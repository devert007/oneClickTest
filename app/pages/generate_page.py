import streamlit as st
from api_utils import upload_document, get_api_response, list_documents, upload_test_pdf,check_document_uniqueness,get_document_text
import uuid
from xml_utils import load_tasks_from_xml, get_tasks,get_preview_tasks
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
from docx import Document
from docx.shared import Pt

def load_css(file_name):
    with open(file_name) as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

load_css("style.css")

try:
    pdfmetrics.registerFont(TTFont('Arial', 'Arial.ttf'))
except:
    st.warning("Не удалось загрузить шрифт Arial. Используется резервный шрифт.")


def markdown_to_word(markdown_text):
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    
    lines = markdown_text.split("\n")
    
    for line in lines:
        while len(line) > 100:
            doc.add_paragraph(line[:100])
            line = line[100:]
        doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
def markdown_to_pdf(markdown_text):
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=letter)

    try:
        p.setFont("Arial", 12)
    except:
        p.setFont("Helvetica", 12)

    lines = markdown_text.split("\n")
    y = 750
    line_height = 20

    for line in lines:
        while len(line) > 100:
            p.drawString(50, y, line[:100])
            line = line[100:]
            y -= line_height
        p.drawString(50, y, line)
        y -= line_height
        
        if y < 50:
            p.showPage()
            y = 750
            try:
                p.setFont("Arial", 12)
            except:
                p.setFont("Helvetica", 12)

    p.save()
    buffer.seek(0)
    return buffer

def show_generate_page():
    st.sidebar.page_link("streamlit_app.py", label="Главная")
    st.sidebar.page_link("pages/generate_page.py", label="Сгенерировать тест")
    st.sidebar.page_link("pages/profile_page.py", label="Профиль")
    st.title("Генерация тестов")

    if 'session_id' not in st.session_state:
        st.session_state.session_id = str(uuid.uuid4())
    if 'is_document_unique' not in st.session_state:
        st.session_state.is_document_unique = False
    if 'uniqueness_checked' not in st.session_state:
        st.session_state.uniqueness_checked = False
    if 'uploaded_file_id' not in st.session_state:
        st.session_state.uploaded_file_id = None
    st.header("1. Выберите или загрузите документ")
    
    documents = list_documents()
    
    if documents:
        doc_options = {doc['filename']: doc['id'] for doc in documents}
        selected_doc = st.selectbox(
            "Выберите из ранее загруженных документов",
            options=[""] + list(doc_options.keys())
        )
        
        if selected_doc:
            st.session_state.uploaded_file_id = doc_options[selected_doc]
            st.info(f"Выбран документ: {selected_doc}")

    uploaded_file = st.file_uploader(
        "Или загрузите новый файл",
        type=["pdf", "docx"],
        key="gen_uploader"
    )
    if uploaded_file: 
        if st.button("Проверить уникальность"):
            with st.spinner("Проверка уникальности..."):
                uniqueness_response = check_document_uniqueness(uploaded_file)
                if uniqueness_response:
                    if uniqueness_response.get("is_unique"):
                        st.success("Документ уникален!")
                        st.session_state.is_document_unique = True
                        st.session_state.uniqueness_checked = True
                    else:
                        st.session_state.is_document_unique = False
                        st.session_state.uniqueness_checked = True
                        source = uniqueness_response.get("source")
                        if source == "SQLite":
                            st.error(f"Документ с именем {uploaded_file.name} уже существует в базе данных!")
                        elif source == "ChromaDB":
                            st.error(
                                f"Документ не уникален! Максимальное сходство: "
                                f"{uniqueness_response['max_similarity']:.2f} с файлом ID {uniqueness_response['similar_doc_id']}"
                            )

        # Show upload button only if document is unique
        if st.session_state.uniqueness_checked and st.session_state.is_document_unique:
            if st.button("Загрузить файл"):

                with st.spinner("Идет загрузка..."):
                    response = upload_document(uploaded_file)
                    if response:
                        st.session_state.uploaded_file_id = response['file_id']
                        st.success(f"Документ {uploaded_file.name} успешно загружен!")
                        st.session_state.uniqueness_checked = False  # Reset for next upload
                        st.session_state.is_document_unique = False
                        documents = list_documents()
                    else:
                        st.error("Ошибка при загрузке документа.")

    

    st.header("2. Настройки теста")
    question_count = st.number_input("Количество вопросов", min_value=1, max_value=20, value=1)
    difficulty = st.select_slider("Уровень сложности", options=["Для начальных классов", "Для средних классов", "Для старших классов", "Для студентов"])
    question_format = st.radio("Формат вопросов", ["Вопрос с выбором ответа", 
    "Вопрос без выбора ответа"])

    st.header("3. Добавить задачи из базы")

    tasks_dict = load_tasks_from_xml()
    subjects = list(tasks_dict.keys()) if tasks_dict else []
    selected_subject = st.selectbox("Выберите предмет", options=[""] + subjects)
    selected_topic = None
    xml_tasks_count = 0

    if selected_subject:
        topics = list(tasks_dict[selected_subject].keys())
        selected_topic = st.selectbox("Выберите тему", options=[""] + topics)
        xml_unlock_tasks_count =get_preview_tasks(subject=selected_subject,
                topic=selected_topic if selected_topic else None,
                difficulty=difficulty,
                task_type=question_format)
        xml_tasks_count = st.number_input("Количество доступных задач по вашим настройкам  из XML", min_value=1 if xml_unlock_tasks_count>0 else 0, max_value=xml_unlock_tasks_count, value=1 if xml_unlock_tasks_count>0 else 0)
    
    st.header("4. Сгенерировать тест")
    if st.button("Создать тест"):
        if 'uploaded_file_id' not in st.session_state and not selected_subject:
            st.error("Выберите документ или предмет для генерации теста!")
            return

        # Adjust question count for AI generation
        ai_question_count = question_count
        generated_questions = []
        answers = ""

        # Generate AI-based questions if needed
        if ai_question_count > 0 and 'uploaded_file_id' in st.session_state:
            print(st.session_state.uploaded_file_id)
            document_text = get_document_text(st.session_state.uploaded_file_id) 
            if question_format == "Вопрос с выбором ответа":
                format_instructions = "Создавай вопросы с 4 вариантами ответов (A, B, C, D). Только один правильный. В конце укажи правильные ответы в формате: 1.A, 2.C и т.д."
            else:
                format_instructions = "Создавай открытые вопросы, требующие развернутого ответа. В конце дай краткие правильные ответы на каждый вопрос."

            # УЛУЧШЕННЫЙ ПРОМПТ
            prompt = f"""
            Сгенерируй качественный тест на основе документа.
            
            ТРЕБОВАНИЯ:
            - Количество вопросов: {ai_question_count}
            - Уровень сложности: {difficulty}
            - Формат вопросов: {question_format}
            - Все вопросы должны быть на русском языке
            - Все вопросы должны быть строго по теме документа
            
            ИНСТРУКЦИИ ПО ФОРМАТУ:
            {format_instructions}
            
            СТРУКТУРА:
            1. Сначала все вопросы по порядку
            2. Затем раздел с правильными ответами
            
            КАЧЕСТВО ВОПРОСОВ:
            - Вопросы должны проверять понимание материала
            - Формулировки четкие и однозначные
            - Охватывай ключевые темы документа
            - Соответствуй уровню сложности: {difficulty}
            
            ВАЖНО:
            - Используй ТОЛЬКО информацию из предоставленного документа
            - Не добавляй вопросы на темы, которых нет в документе
            - Все ответы должны быть подтверждены текстом документа
            
            Текст документа:
            {document_text}
            
            Сгенерируй тест строго по этим требованиям.
            """
            
            with st.spinner("Генерация качественного теста..."):
                response = get_api_response(
                    question=prompt,
                    session_id=st.session_state.session_id,
                    model="smol-lm-3b" # Меняем на русскую модель
                )
                
                if response:
                    test_content = response['answer']

                    # Сохраняем весь тест как есть
                    st.session_state.generated_test = test_content
                    st.session_state.test_generated = True
                    st.session_state.test_pdf = markdown_to_pdf(test_content)
                    st.session_state.test_word = markdown_to_word(test_content)
                    st.success("✅ Тест успешно сгенерирован!")


        # Остальной код с XML tasks остается без изменений...
        # Add XML tasks
        xml_questions = ""
        print(selected_subject,xml_tasks_count)

        if selected_subject and xml_tasks_count > 0:
            print('tyta')
            xml_tasks = get_tasks(
                subject=selected_subject,
                topic=selected_topic if selected_topic else None,
                difficulty=difficulty,
                task_type=question_format,
                limit=xml_tasks_count
            )
            print(xml_tasks)
            for i, task in enumerate(xml_tasks, len(generated_questions) + 1):
                xml_questions.append(f"{i}. {task.question}")
                answers.append(f"[{i}].{task.answer}")
        print(test_content)
        # Combine questions and answers
        combined_questions = test_content + xml_questions
        if not combined_questions:
            st.error("Не удалось сгенерировать или выбрать вопросы!")
            return

        test_content = "".join(combined_questions + answers)
        test_content=test_content.split("</think>")[-1]
        st.session_state.generated_test = test_content
        st.session_state.test_generated = True
        st.session_state.test_pdf = markdown_to_pdf(test_content)
        st.session_state.test_word = markdown_to_word(test_content)
        st.success("Тест успешно сгенерирован!")

        pdf_response = upload_test_pdf(
            pdf_buffer=st.session_state.test_pdf,
            filename="generated_test.pdf",
            document_id=st.session_state.uploaded_file_id if 'uploaded_file_id' in st.session_state else None,
            session_id=st.session_state.session_id
        )
        if pdf_response:
            st.success(f"PDF тест сохранён! ID: {pdf_response.get('file_id')}")
        else:
            st.error("Ошибка сохранения PDF теста")

    if 'test_generated' in st.session_state and st.session_state.test_generated and 'generated_test' in st.session_state:
        with st.expander("Просмотр теста", expanded=True):
            st.markdown(st.session_state.generated_test)

        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.download_button(
                label="Скачать MD",
                data=st.session_state.generated_test,
                file_name="generated_test.md",
                mime="text/markdown",
                key="download_md"
            )
        
        with col2:
            st.download_button(
                label="Скачать PDF",
                data=st.session_state.test_pdf,
                file_name="generated_test.pdf",
                mime="application/pdf",
                key="download_pdf"
            )
        with col3:
            st.download_button(
                label="Скачать WORD",
                data=st.session_state.test_word,
                file_name="generated_test.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_word"
            )



if __name__ == "__main__":
    show_generate_page()